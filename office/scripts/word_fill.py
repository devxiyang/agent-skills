#!/usr/bin/env python3
"""
Fill a Word template by replacing {{placeholders}} with values from a JSON file.

Usage:
    python3 word_fill.py template.docx data.json
    python3 word_fill.py template.docx data.json --output filled.docx
    python3 word_fill.py template.docx --set name=Alice --set date=2025-01-01

Template format:
    Use {{key}} anywhere in the document — paragraphs, tables, headers, footers.
    Example: "Dear {{name}}, your invoice for {{amount}} is due on {{due_date}}."

JSON format:
    {"name": "Alice", "amount": "$1,200", "due_date": "2025-02-01"}
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print('Missing dependency. Run: pip install python-docx')
    sys.exit(1)


def replace_in_paragraph(para, replacements: dict[str, str]) -> bool:
    """Replace placeholders in a paragraph, preserving runs as much as possible."""
    full_text = ''.join(run.text for run in para.runs)
    changed = False
    for key, value in replacements.items():
        placeholder = f'{{{{{key}}}}}'
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))
            changed = True
    if changed:
        # Clear all runs and put text in first run to preserve style
        if para.runs:
            para.runs[0].text = full_text
            for run in para.runs[1:]:
                run.text = ''
    return changed


def fill_document(doc: 'Document', replacements: dict[str, str]) -> int:
    count = 0

    # Paragraphs
    for para in doc.paragraphs:
        if replace_in_paragraph(para, replacements):
            count += 1

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_paragraph(para, replacements):
                        count += 1

    # Headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            if replace_in_paragraph(para, replacements):
                count += 1
        for para in section.footer.paragraphs:
            if replace_in_paragraph(para, replacements):
                count += 1

    return count


def main() -> None:
    parser = argparse.ArgumentParser(description='Fill a Word template with data')
    parser.add_argument('template', help='Word template (.docx) with {{placeholders}}')
    parser.add_argument('data', nargs='?', help='JSON file with replacement values')
    parser.add_argument('--output', '-o', help='Output .docx (default: <template>_filled.docx)')
    parser.add_argument('--set', action='append', metavar='KEY=VALUE',
                        help='Set a value directly (can be used multiple times)')
    args = parser.parse_args()

    tmpl = Path(args.template)
    if not tmpl.exists():
        print(f'Error: {tmpl} not found')
        sys.exit(1)

    replacements: dict[str, str] = {}

    if args.data:
        data_path = Path(args.data)
        if not data_path.exists():
            print(f'Error: {data_path} not found')
            sys.exit(1)
        replacements.update(json.loads(data_path.read_text()))

    if args.set:
        for item in args.set:
            if '=' not in item:
                print(f'Error: --set must be KEY=VALUE, got: {item}')
                sys.exit(1)
            k, v = item.split('=', 1)
            replacements[k.strip()] = v.strip()

    if not replacements:
        print('Error: provide a JSON file or --set KEY=VALUE')
        sys.exit(1)

    out = Path(args.output) if args.output else tmpl.with_name(tmpl.stem + '_filled.docx')

    print(f'Loading {tmpl}...')
    doc = Document(str(tmpl))

    print(f'Replacing {len(replacements)} placeholder(s)...')
    count = fill_document(doc, replacements)

    doc.save(str(out))
    print(f'Saved → {out} ({count} paragraph(s) updated)')


if __name__ == '__main__':
    main()
